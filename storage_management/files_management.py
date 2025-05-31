import os
import fitz  # PyMuPDF
import docx
import tempfile
import msgpack
import storage_management.utils as utils


class TextSector:
    def __init__(self, filename, text, start=0, end=0, index_id=-1):
        self.filename = filename
        self.text = text
        self.start = start
        self.end = start + len(text) if text else end
        self.index_id = index_id

    def to_dict(self):
        return {
            'filename': self.filename,
            'start': self.start,
            'end': self.end,
            'index_id': self.index_id
        }

    @classmethod
    def from_dict(cls, data):
        return cls(
            filename=data['filename'],
            text="",
            start=data['start'],
            end=data['end'],
            index_id=data['index_id']
        )


def serialize_text_sectors(file_sectors: list[TextSector], file_path: str):
    dict_list = [sector.to_dict() for sector in file_sectors]
    packed = msgpack.packb(dict_list, use_bin_type=True)
    with open(file_path, 'wb') as f:
        f.write(packed)


def deserialize_text_sectors(file_path: str) -> list[TextSector]:
    with open(file_path, 'rb') as f:
        dict_list = msgpack.unpackb(f.read(), raw=False)
        return [TextSector.from_dict(d) for d in dict_list]


class Reader:
    def __init__(self):
        self.filename = ""
        self.file_handler = None
        self.content = ""
        self.texts = []

    def read_index(self, file, knowledge_base, default_index_name, root_dir):
        if not file.name.endswith(".faiss"):
            index_name = default_index_name
        index_name = file.name
        descriptor = index_name.split(".")[0] + ".msgpack"
        paths = utils.get_paths_from_names({index_name, descriptor}, root_dir)

        knowledge_base.read_index_file(paths[index_name])
        return paths[index_name], paths[descriptor], deserialize_text_sectors(paths[descriptor])

    def read_file(self, file, isPath=False) -> str:
        if isPath:
            ext = file.split(".")[-1]
            if ext == "txt":
                self.content = self.read_txt(file, isPath)
            elif ext == "pdf":
                self.content = self.read_pdf(file, isPath)
            elif ext == "docx":
                self.content = self.read_dock(file, isPath)
        else:
            if file.name.endswith(".txt"):
                self.content = self.read_txt(file)
            elif file.name.endswith(".pdf"):
                self.content = self.read_pdf(file)
            elif file.name.endswith(".docx"):
                self.content = self.read_dock(file)

        return self.content

    def read_txt(self, file, isPath=False) -> str:
        if isPath:
            with open(file, 'r', encoding="utf-8") as f:
                return f.read()
        else:
            return file.read().decode("utf-8")

    def read_pdf(self, file, isPath=False) -> str:
        if isPath:
            doc = fitz.open(file)
            buf = [page.get_text() for page in doc]
            doc.close()
        else:
            with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
                tmp_file.write(file.read())
                tmp_file_path = tmp_file.name

            doc = fitz.open(tmp_file_path)
            buf = [page.get_text() for page in doc]
            doc.close()
            os.unlink(tmp_file_path)
        return '\n'.join(buf)

    def read_dock(self, file, isPath=False) -> str:
        if isPath:
            doc = docx.Document(file)
        else:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                tmp_file.write(file.read())
                tmp_file_path = tmp_file.name
            doc = docx.Document(tmp_file_path)
            os.unlink(tmp_file_path)

        buf = [para.text for para in doc.paragraphs if para.text.strip() != ""]
        return '\n'.join(buf)



